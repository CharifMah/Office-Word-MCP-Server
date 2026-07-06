"""Content creation tools - add headings, paragraphs, tables, images, page breaks, bullet lists."""

import os
from typing import List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def add_heading(filename: str, text: str, level: int = 1,
                      font_name: Optional[str] = None, font_size: Optional[int] = None,
                      bold: Optional[bool] = None, italic: Optional[bool] = None,
                      border_bottom: bool = False,
                      page_break_before: bool = False,
                      keep_with_next: bool = False) -> str:
    """Add a heading to a Word document with optional formatting."""
    filename = ensure_docx_extension(filename)
    try:
        level = int(level)
    except (ValueError, TypeError):
        return "Invalid parameter: level must be an integer between 1 and 9"
    if level < 1 or level > 9:
        return f"Invalid heading level: {level}. Level must be between 1 and 9."
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    try:
        doc = Document(filename)
        ensure_heading_style(doc)
        try:
            heading = doc.add_heading(text, level=level)
        except Exception:
            heading = doc.add_paragraph(text)
            heading.style = doc.styles['Normal']
            if heading.runs:
                run = heading.runs[0]
                run.bold = True
                run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
        if any([font_name, font_size, bold is not None, italic is not None]):
            for run in heading.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
        pPr = heading._element.get_or_add_pPr()
        if border_bottom:
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)
        if page_break_before:
            pb = OxmlElement('w:pageBreakBefore')
            pb.set(qn('w:val'), 'true')
            pPr.append(pb)
        if keep_with_next:
            kwn = OxmlElement('w:keepNext')
            kwn.set(qn('w:val'), 'true')
            pPr.append(kwn)
            kwn2 = OxmlElement('w:keepLines')
            kwn2.set(qn('w:val'), 'true')
            pPr.append(kwn2)
        doc.save(filename)
        return f"Heading '{text}' (level {level}) added to {filename}"
    except Exception as e:
        return f"Failed to add heading: {str(e)}"


async def add_paragraph(filename: str, text: str, style: Optional[str] = None,
                        font_name: Optional[str] = None, font_size: Optional[int] = None,
                        bold: Optional[bool] = None, italic: Optional[bool] = None,
                        color: Optional[str] = None) -> str:
    """Add a paragraph to a Word document with optional formatting."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    try:
        doc = Document(filename)
        para = doc.add_paragraph(text, style=style)
        if any([font_name, font_size, bold is not None, italic is not None, color]):
            for run in para.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                if color:
                    color_hex = color.lstrip('#')
                    run.font.color.rgb = RGBColor.from_string(color_hex)
        doc.save(filename)
        return f"Paragraph added to {filename}"
    except Exception as e:
        return f"Failed to add paragraph: {str(e)}"


async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        ensure_table_style(doc)
        if data:
            for i, row_data in enumerate(data):
                if i < rows:
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            table.cell(i, j).text = str(cell_text)
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"


async def add_picture(filename: str, image_path: str, width: Optional[float] = None) -> str:
    """Add a picture to a Word document."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    if not os.path.exists(image_path):
        return f"Image file {image_path} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    try:
        doc = Document(filename)
        if width:
            doc.add_picture(image_path, width=Inches(width))
        else:
            doc.add_picture(image_path)
        doc.save(filename)
        return f"Picture '{image_path}' added to {filename}"
    except Exception as e:
        return f"Failed to add picture: {str(e)}"


async def add_page_break(filename: str) -> str:
    """Add a page break to a Word document."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    try:
        doc = Document(filename)
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}"
    except Exception as e:
        return f"Failed to add page break: {str(e)}"


async def add_bullet_list(filename: str, items: List[str], style: str = 'List Bullet') -> str:
    """Add a bulleted list to the end of the document."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        for item in items:
            doc.add_paragraph(item, style=style)
        doc.save(filename)
        return f"Added {len(items)} bullet items"
    except Exception as e:
        return f"Failed to add bullet list: {str(e)}"