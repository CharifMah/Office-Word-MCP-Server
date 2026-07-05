"""Content modification tools - delete, search/replace, insert at index, move sections, spacing, header/footer."""

import os
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension


async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph by its index."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index {paragraph_index} out of range (0-{len(doc.paragraphs)-1})"
        para = doc.paragraphs[paragraph_index]
        para._element.getparent().remove(para._element)
        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully"
    except Exception as e:
        return f"Failed to delete paragraph: {str(e)}"


async def search_and_replace(filename: str, find_text: str, replace_text: str) -> str:
    """Search for text and replace all occurrences."""
    from word_document_server.utils.document_utils import find_and_replace_text
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        count = find_and_replace_text(filename, find_text, replace_text)
        return f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'"
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"


async def insert_paragraph_at_index(filename: str, text: str, paragraph_index: int = 0,
                                     style: Optional[str] = None,
                                     font_name: Optional[str] = None, font_size: Optional[int] = None,
                                     bold: Optional[bool] = None, italic: Optional[bool] = None,
                                     color: Optional[str] = None) -> str:
    """Insert a paragraph at a specific index in the document."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        paragraphs = doc.paragraphs
        if paragraph_index >= len(paragraphs):
            para = doc.add_paragraph(text)
        else:
            target_para = paragraphs[paragraph_index]
            new_p = OxmlElement('w:p')
            target_para._element.addprevious(new_p)
            para = doc.paragraphs[paragraph_index]
        para.text = text
        if style:
            try:
                para.style = doc.styles[style]
            except KeyError:
                pass
        if any([font_name, font_size, bold is not None, italic is not None, color]):
            for run in para.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                if color:
                    color_hex = color.lstrip('#')
                    run.font.color.rgb = RGBColor.from_string(color_hex)
        doc.save(filename)
        return f"Paragraph inserted at index {paragraph_index}"
    except Exception as e:
        return f"Failed to insert paragraph: {str(e)}"


async def move_section(filename: str, start_heading_text: str, target_heading_text: str) -> str:
    """Move a section (from start_heading to next heading of same level) to before target_heading."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        # Find the start heading
        start_idx = None
        start_level = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == start_heading_text and para.style.name.startswith('Heading'):
                start_idx = i
                start_level = para.style.name
                break
        if start_idx is None:
            return f"Heading '{start_heading_text}' not found"

        # Find the end of the section (next heading of same level)
        end_idx = len(doc.paragraphs)
        for i in range(start_idx + 1, len(doc.paragraphs)):
            if doc.paragraphs[i].style.name == start_level:
                end_idx = i
                break

        # Find the target heading
        target_idx = None
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == target_heading_text and para.style.name.startswith('Heading'):
                target_idx = i
                break
        if target_idx is None:
            return f"Target heading '{target_heading_text}' not found"

        # Extract the section elements
        section_elements = []
        for i in range(start_idx, end_idx):
            section_elements.append(doc.paragraphs[i]._element)

        # Insert before target
        target_element = doc.paragraphs[target_idx]._element
        for elem in section_elements:
            target_element.addprevious(elem)

        doc.save(filename)
        return f"Section '{start_heading_text}' moved before '{target_heading_text}'"
    except Exception as e:
        return f"Failed to move section: {str(e)}"


async def delete_paragraphs_range(filename: str, start_index: int, end_index: int) -> str:
    """Delete a range of paragraphs from start_index to end_index (inclusive)."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        if start_index < 0 or end_index >= len(doc.paragraphs) or start_index > end_index:
            return f"Invalid range: {start_index}-{end_index}"
        count = 0
        for i in range(end_index, start_index - 1, -1):
            para = doc.paragraphs[i]
            para._element.getparent().remove(para._element)
            count += 1
        doc.save(filename)
        return f"Deleted {count} paragraphs (indices {start_index}-{end_index})"
    except Exception as e:
        return f"Failed to delete paragraphs: {str(e)}"


async def set_paragraph_spacing(filename: str, paragraph_index: int,
                                 before: Optional[float] = None,
                                 after: Optional[float] = None,
                                 line_spacing: Optional[float] = None) -> str:
    """Set spacing for a specific paragraph."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index {paragraph_index} out of range"
        para = doc.paragraphs[paragraph_index]
        pf = para.paragraph_format
        if before is not None:
            pf.space_before = Pt(before)
        if after is not None:
            pf.space_after = Pt(after)
        if line_spacing is not None:
            pf.line_spacing = line_spacing
        doc.save(filename)
        return f"Spacing set for paragraph {paragraph_index}"
    except Exception as e:
        return f"Failed to set spacing: {str(e)}"


async def set_header_footer(filename: str, header_text: str = None,
                             footer_text: str = None) -> str:
    """Set header and/or footer text for the document."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        for section in doc.sections:
            if header_text is not None:
                header = section.header
                if header.paragraphs:
                    header.paragraphs[0].text = header_text
                else:
                    header.add_paragraph(header_text)
            if footer_text is not None:
                footer = section.footer
                if footer.paragraphs:
                    footer.paragraphs[0].text = footer_text
                else:
                    footer.add_paragraph(footer_text)
        doc.save(filename)
        return f"Header/footer set"
    except Exception as e:
        return f"Failed to set header/footer: {str(e)}"


async def add_page_number_to_footer(filename: str) -> str:
    """Add page number field to the footer of the document."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        for section in doc.sections:
            footer = section.footer
            if footer.paragraphs:
                para = footer.paragraphs[0]
            else:
                para = footer.add_paragraph()
            run = para.add_run()
            fld_char1 = OxmlElement('w:fldChar')
            fld_char1.set(qn('w:fldCharType'), 'begin')
            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            instr_text.text = 'PAGE'
            fld_char2 = OxmlElement('w:fldChar')
            fld_char2.set(qn('w:fldCharType'), 'end')
            run._element.append(fld_char1)
            run._element.append(instr_text)
            run._element.append(fld_char2)
        doc.save(filename)
        return f"Page number added to footer"
    except Exception as e:
        return f"Failed to add page number: {str(e)}"