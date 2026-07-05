"""Read tools - get paragraph text by index."""

import os
from docx import Document

from word_document_server.utils.file_utils import ensure_docx_extension


async def get_paragraph_text(filename: str, paragraph_index: int) -> str:
    """Get the text of a specific paragraph by index."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    try:
        doc = Document(filename)
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index {paragraph_index} out of range (0-{len(doc.paragraphs)-1})"
        para = doc.paragraphs[paragraph_index]
        return f"Paragraph {paragraph_index} [{para.style.name}]: {para.text}"
    except Exception as e:
        return f"Failed to get paragraph: {str(e)}"