"""Table tools - delete, insert at position, format all cells, replace data, borders, margins, indentation."""

import os
from typing import List, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension


async def delete_table(filename: str, table_index: int) -> str:
    """Delete a table from the document by its index (0-based)."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index {table_index} out of range (0-{len(doc.tables)-1})"
        table = doc.tables[table_index]
        table._element.getparent().remove(table._element)
        doc.save(filename)
        return f"Table {table_index} deleted from document"
    except Exception as e:
        return f"Failed to delete table: {str(e)}"


async def insert_table_at_position(filename: str, headers: List[str], data: List[List[str]],
                                    target_text: str = None, target_paragraph_index: int = None,
                                    position: str = 'after',
                                    header_color: str = 'E97132',
                                    header_text_color: str = 'FFFFFF',
                                    font_size: int = 9,
                                    horizontal_align: str = 'center',
                                    vertical_align: str = 'center',
                                    border_style: str = 'single',
                                    border_size: int = 4,
                                    border_color: str = '000000') -> str:
    """Insert a clean formatted table at a specific position in the document.

    The table is inserted directly at the target paragraph (before/after) instead of
    being created at the end and moved. All cells are centered horizontally and vertically.
    Header uses the requested colors and text formatting.
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        total_rows = len(data) + 1
        total_cols = len(headers)

        # Find anchor paragraph
        target_para = None
        if target_paragraph_index is not None:
            if 0 <= target_paragraph_index < len(doc.paragraphs):
                target_para = doc.paragraphs[target_paragraph_index]
        elif target_text:
            for para in doc.paragraphs:
                if target_text in para.text:
                    target_para = para
                    break

        if target_para is None:
            return "Target paragraph not found"

        # Create table element directly at anchor to avoid trailing blank paragraphs
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        p = target_para._element
        new_tbl = OxmlElement('w:tbl')
        if position == 'before':
            p.addprevious(new_tbl)
        else:
            p.addnext(new_tbl)
        table = doc.tables[-1]  # newest table element maps to last Table object

        # Ensure table has the requested number of rows/columns
        while len(table.rows) < total_rows:
            table.add_row()
        while len(table.columns) < total_cols:
            table.add_column()

        # Helper to set cell text without extra leading space
        def _set_cell_text(cell, text):
            cell.text = ''
            p = cell.paragraphs[0]
            p.text = str(text) if text is not None else ''
            return p

        align_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para_align = align_map.get(horizontal_align.lower(), WD_ALIGN_PARAGRAPH.CENTER)

        # Fill headers
        for c, h in enumerate(headers):
            if c < total_cols:
                cell = table.cell(0, c)
                para = _set_cell_text(cell, h)
                para.alignment = para_align
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(font_size)
                    if header_text_color:
                        run.font.color.rgb = RGBColor.from_string(header_text_color.lstrip('#'))
                # header background
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), header_color.lstrip('#'))
                tc_pr.append(shd)
                # vertical align
                valign = OxmlElement('w:vAlign')
                valign.set(qn('w:val'), vertical_align.lower())
                tc_pr.append(valign)

        # Fill data rows
        for r, row_data in enumerate(data):
            for c, val in enumerate(row_data):
                if c < total_cols:
                    cell = table.cell(r + 1, c)
                    para = _set_cell_text(cell, val)
                    para.alignment = para_align
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.size = Pt(font_size)
                    tc_pr = cell._tc.get_or_add_tcPr()
                    valign = OxmlElement('w:vAlign')
                    valign.set(qn('w:val'), vertical_align.lower())
                    tc_pr.append(valign)

        # Set borders
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._element.insert(0, tblPr)
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        borders = OxmlElement('w:tblBorders')
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), border_style)
            border.set(qn('w:sz'), str(border_size))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            borders.append(border)
        tblPr.append(borders)

        doc.save(filename)
        return f"Table with {len(data)} rows inserted at position"
    except Exception as e:
        return f"Failed to insert table: {str(e)}"


async def format_table_all_cells(filename: str, table_index: int,
                                  font_size: Optional[int] = None,
                                  bold: Optional[bool] = None,
                                  header_bold: Optional[bool] = None,
                                  header_color: Optional[str] = None,
                                  header_text_color: Optional[str] = None,
                                  align: Optional[str] = None,
                                  style_name: Optional[str] = None) -> str:
    """Format all cells of a table in one call."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index {table_index} out of range"
        table = doc.tables[table_index]
        if style_name:
            try:
                table.style = doc.styles[style_name]
            except KeyError:
                pass
        align_map = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER, 'right': WD_ALIGN_PARAGRAPH.RIGHT}
        para_align = align_map.get(align.lower()) if align else None
        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                cell = table.cell(r, c)
                for para in cell.paragraphs:
                    if para_align is not None:
                        para.alignment = para_align
                    pf = para.paragraph_format
                    pf.left_indent = Cm(0)
                    pf.first_line_indent = Cm(0)
                    for run in para.runs:
                        if font_size is not None:
                            run.font.size = Pt(font_size)
                        if bold is not None:
                            run.font.bold = bold
            if r == 0:
                if header_bold is not None:
                    for c in range(len(table.columns)):
                        for para in table.cell(0, c).paragraphs:
                            for run in para.runs:
                                run.font.bold = header_bold
                if header_color:
                    for c in range(len(table.columns)):
                        cell = table.cell(0, c)
                        tc_pr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), header_color)
                        tc_pr.append(shd)
                if header_text_color:
                    color_hex = header_text_color.lstrip('#')
                    for c in range(len(table.columns)):
                        for para in table.cell(0, c).paragraphs:
                            for run in para.runs:
                                run.font.color.rgb = RGBColor.from_string(color_hex)
        doc.save(filename)
        return f"All cells in table {table_index} formatted successfully"
    except Exception as e:
        return f"Failed to format table: {str(e)}"


async def replace_table_data(filename: str, table_index: int, headers: List[str], data: List[List[str]]) -> str:
    """Replace all data in a table with new headers and data."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index {table_index} out of range"
        table = doc.tables[table_index]
        total_rows = len(data) + 1
        while len(table.rows) < total_rows:
            table.add_row()
        while len(table.rows) > total_rows:
            row = table.rows[-1]
            row._element.getparent().remove(row._element)
        for c, h in enumerate(headers):
            if c < len(table.columns):
                cell = table.cell(0, c)
                cell.text = h
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
        for r, row_data in enumerate(data):
            for c, val in enumerate(row_data):
                if c < len(table.columns):
                    cell = table.cell(r + 1, c)
                    cell.text = str(val) if val else ''
        doc.save(filename)
        return f"Table {table_index} data replaced"
    except Exception as e:
        return f"Failed to replace table data: {str(e)}"


async def set_table_borders(filename: str, table_index: int,
                             border_style: str = 'single', border_size: int = 4,
                             border_color: str = '000000') -> str:
    """Set borders for a table."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index {table_index} out of range"
        table = doc.tables[table_index]
        tblPr = table._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._element.insert(0, tblPr)
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        borders = OxmlElement('w:tblBorders')
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{edge}')
            border.set(qn('w:val'), border_style)
            border.set(qn('w:sz'), str(border_size))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            borders.append(border)
        tblPr.append(borders)
        doc.save(filename)
        return f"Table {table_index} borders set"
    except Exception as e:
        return f"Failed to set borders: {str(e)}"


async def set_cell_margins_all(filename: str, table_index: int,
                                top: float = 0, bottom: float = 0,
                                left: float = 0.19, right: float = 0.19) -> str:
    """Set cell margins (padding) for all cells in a table (in cm)."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index {table_index} out of range"
        table = doc.tables[table_index]
        top_tw = int(top * 567)
        bottom_tw = int(bottom * 567)
        left_tw = int(left * 567)
        right_tw = int(right * 567)
        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                cell = table.cell(r, c)
                tcPr = cell._element.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    cell._element.insert(0, tcPr)
                existing = tcPr.find(qn('w:tcMar'))
                if existing is not None:
                    tcPr.remove(existing)
                tcMar = OxmlElement('w:tcMar')
                for edge, val in [('top', top_tw), ('bottom', bottom_tw), ('left', left_tw), ('right', right_tw)]:
                    margin = OxmlElement(f'w:{edge}')
                    margin.set(qn('w:w'), str(val))
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)
        doc.save(filename)
        return f"Cell margins set for table {table_index}"
    except Exception as e:
        return f"Failed to set cell margins: {str(e)}"


async def remove_all_indentation_from_table(filename: str, table_index: int) -> str:
    """Remove all paragraph indentation from all cells in a table."""
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}"
    try:
        doc = Document(filename)
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index {table_index} out of range"
        table = doc.tables[table_index]
        count = 0
        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                cell = table.cell(r, c)
                for para in cell.paragraphs:
                    pf = para.paragraph_format
                    pf.left_indent = Cm(0)
                    pf.first_line_indent = Cm(0)
                    pf.right_indent = Cm(0)
                    pPr = para._element.find(qn('w:pPr'))
                    if pPr is not None:
                        ind = pPr.find(qn('w:ind'))
                        if ind is not None:
                            pPr.remove(ind)
                            count += 1
        doc.save(filename)
        return f"Removed indentation from {count} paragraphs in table {table_index}"
    except Exception as e:
        return f"Failed to remove indentation: {str(e)}"