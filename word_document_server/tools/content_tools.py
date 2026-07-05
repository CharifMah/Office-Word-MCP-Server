"""
Content tools for Word Document Server.

These tools add various types of content to Word documents,
including headings, paragraphs, tables, images, and page breaks.
"""
import os
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.utils.document_utils import find_and_replace_text, insert_header_near_text, insert_numbered_list_near_text, insert_line_or_paragraph_near_text, replace_paragraph_block_below_header, replace_block_between_manual_anchors
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def add_heading(filename: str, text: str, level: int = 1,
                      font_name: Optional[str] = None, font_size: Optional[int] = None,
                      bold: Optional[bool] = None, italic: Optional[bool] = None,
                      border_bottom: bool = False) -> str:
    """Add a heading to a Word document with optional formatting.

    Args:
        filename: Path to the Word document
        text: Heading text
        level: Heading level (1-9, where 1 is the highest level)
        font_name: Font family (e.g., 'Helvetica')
        font_size: Font size in points (e.g., 14)
        bold: True/False for bold text
        italic: True/False for italic text
        border_bottom: True to add bottom border (for section headers)
    """
    filename = ensure_docx_extension(filename)

    # Ensure level is converted to integer
    try:
        level = int(level)
    except (ValueError, TypeError):
        return "Invalid parameter: level must be an integer between 1 and 9"

    # Validate level range
    if level < 1 or level > 9:
        return f"Invalid heading level: {level}. Level must be between 1 and 9."

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)

        # Ensure heading styles exist
        ensure_heading_style(doc)

        # Try to add heading with style
        try:
            heading = doc.add_heading(text, level=level)
        except Exception as style_error:
            # If style-based approach fails, use direct formatting
            heading = doc.add_paragraph(text)
            heading.style = doc.styles['Normal']
            if heading.runs:
                run = heading.runs[0]
                run.bold = True
                # Adjust size based on heading level
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)

        # Apply formatting to all runs in the heading
        if any([font_name, font_size, bold is not None, italic is not None]):
            for run in heading.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic

        # Add bottom border if requested
        if border_bottom:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            pPr = heading._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')  # 0.5pt border
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')

            pBdr.append(bottom)
            pPr.append(pBdr)

        doc.save(filename)
        return f"Heading '{text}' (level {level}) added to {filename}"
    except Exception as e:
        return f"Failed to add heading: {str(e)}"


async def add_paragraph(filename: str, text: str, style: Optional[str] = None,
                        font_name: Optional[str] = None, font_size: Optional[int] = None,
                        bold: Optional[bool] = None, italic: Optional[bool] = None,
                        color: Optional[str] = None) -> str:
    """Add a paragraph to a Word document with optional formatting.

    Args:
        filename: Path to the Word document
        text: Paragraph text
        style: Optional paragraph style name
        font_name: Font family (e.g., 'Helvetica', 'Times New Roman')
        font_size: Font size in points (e.g., 14, 36)
        bold: True/False for bold text
        italic: True/False for italic text
        color: RGB color as hex string (e.g., '000000' for black)
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)
        paragraph = doc.add_paragraph(text)

        if style:
            try:
                paragraph.style = style
            except KeyError:
                # Style doesn't exist, use normal and report it
                paragraph.style = doc.styles['Normal']
                doc.save(filename)
                return f"Style '{style}' not found, paragraph added with default style to {filename}"

        # Apply formatting to all runs in the paragraph
        if any([font_name, font_size, bold is not None, italic is not None, color]):
            for run in paragraph.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                if color:
                    # Remove any '#' prefix if present
                    color_hex = color.lstrip('#')
                    run.font.color.rgb = RGBColor.from_string(color_hex)

        doc.save(filename)
        return f"Paragraph added to {filename}"
    except Exception as e:
        return f"Failed to add paragraph: {str(e)}"


async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        filename: Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        
        # Try to set the table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If style doesn't exist, add basic borders
            pass
        
        # Fill table with data if provided
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"


async def add_picture(filename: str, image_path: str, width: Optional[float] = None) -> str:
    """Add an image to a Word document.
    
    Args:
        filename: Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    filename = ensure_docx_extension(filename)
    
    # Validate document existence
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)
    
    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        return f"Image file not found: {abs_image_path}"
    
    # Check image file size
    try:
        image_size = os.path.getsize(abs_image_path) / 1024  # Size in KB
        if image_size <= 0:
            return f"Image file appears to be empty: {abs_image_path} (0 KB)"
    except Exception as size_error:
        return f"Error checking image file: {str(size_error)}"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(abs_filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(abs_filename)
        # Additional diagnostic info
        diagnostic = f"Attempting to add image ({abs_image_path}, {image_size:.2f} KB) to document ({abs_filename})"
        
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            # More detailed error for the specific operation
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            return f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}"
    except Exception as outer_error:
        # Fallback error handling
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        return f"Document processing error: {error_type} - {error_msg or 'No error details available'}"


async def add_page_break(filename: str) -> str:
    """Add a page break to the document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}."
    except Exception as e:
        return f"Failed to add page break: {str(e)}"


async def add_table_of_contents(filename: str, title: str = "Table of Contents", max_level: int = 3) -> str:
    """Add a table of contents to a Word document based on heading styles.
    
    Args:
        filename: Path to the Word document
        title: Optional title for the table of contents
        max_level: Maximum heading level to include (1-9)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        # Ensure max_level is within valid range
        max_level = max(1, min(max_level, 9))
        
        doc = Document(filename)
        
        # Collect headings and their positions
        headings = []
        for i, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph style is a heading
            if paragraph.style and paragraph.style.name.startswith('Heading '):
                try:
                    # Extract heading level from style name
                    level = int(paragraph.style.name.split(' ')[1])
                    if level <= max_level:
                        headings.append({
                            'level': level,
                            'text': paragraph.text,
                            'position': i
                        })
                except (ValueError, IndexError):
                    # Skip if heading level can't be determined
                    pass
        
        if not headings:
            return f"No headings found in document {filename}. Table of contents not created."
        
        # Create a new document with the TOC
        toc_doc = Document()
        
        # Add title
        if title:
            toc_doc.add_heading(title, level=1)
        
        # Add TOC entries
        for heading in headings:
            # Indent based on level (using tab characters)
            indent = '    ' * (heading['level'] - 1)
            toc_doc.add_paragraph(f"{indent}{heading['text']}")
        
        # Add page break
        toc_doc.add_page_break()
        
        # Get content from original document
        for paragraph in doc.paragraphs:
            p = toc_doc.add_paragraph(paragraph.text)
            # Copy style if possible
            try:
                if paragraph.style:
                    p.style = paragraph.style.name
            except:
                pass
        
        # Copy tables
        for table in doc.tables:
            # Create a new table with the same dimensions
            new_table = toc_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            # Copy cell contents
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        new_table.cell(i, j).text = paragraph.text
        
        # Save the new document with TOC
        toc_doc.save(filename)
        
        return f"Table of contents with {len(headings)} entries added to {filename}"
    except Exception as e:
        return f"Failed to add table of contents: {str(e)}"


async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph from a document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to delete (0-based)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Delete the paragraph (by removing its content and setting it empty)
        # Note: python-docx doesn't support true paragraph deletion, this is a workaround
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._p
        p.getparent().remove(p)
        
        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully."
    except Exception as e:
        return f"Failed to delete paragraph: {str(e)}"


async def search_and_replace(filename: str, find_text: str, replace_text: str) -> str:
    """Search for text and replace all occurrences.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Perform find and replace
        count = find_and_replace_text(doc, find_text, replace_text)
        
        if count > 0:
            doc.save(filename)
            return f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'."
        else:
            return f"No occurrences of '{find_text}' found."
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"

async def insert_header_near_text_tool(filename: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_header_near_text(filename, target_text, header_title, position, header_style, target_paragraph_index)

async def insert_numbered_list_near_text_tool(filename: str, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None, bullet_type: str = 'bullet') -> str:
    """Insert a bulleted or numbered list before or after the target paragraph. Specify by text or paragraph index."""
    return insert_numbered_list_near_text(filename, target_text, list_items, position, target_paragraph_index, bullet_type)

async def insert_line_or_paragraph_near_text_tool(filename: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None) -> str:
    """Insert a new line or paragraph (with specified or matched style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_line_or_paragraph_near_text(filename, target_text, line_text, position, line_style, target_paragraph_index)

async def replace_paragraph_block_below_header_tool(filename: str, header_text: str, new_paragraphs: list, detect_block_end_fn=None) -> str:
    """Reemplaza el bloque de párrafos debajo de un encabezado, evitando modificar TOC."""
    return replace_paragraph_block_below_header(filename, header_text, new_paragraphs, detect_block_end_fn)

async def replace_block_between_manual_anchors_tool(filename: str, start_anchor_text: str, new_paragraphs: list, end_anchor_text: str = None, match_fn=None, new_paragraph_style: str = None) -> str:
    """Replace all content between start_anchor_text and end_anchor_text (or next logical header if not provided)."""
    return replace_block_between_manual_anchors(filename, start_anchor_text, new_paragraphs, end_anchor_text, match_fn, new_paragraph_style)


async def insert_paragraph_at_index(filename: str, text: str, paragraph_index: int = 0, style: Optional[str] = None,
                                     font_name: Optional[str] = None, font_size: Optional[int] = None,
                                     bold: Optional[bool] = None, italic: Optional[bool] = None,
                                     color: Optional[str] = None) -> str:
    """Insert a paragraph at a specific index in the document.

    Args:
        filename: Path to the Word document
        text: Paragraph text
        paragraph_index: Index at which to insert the paragraph (0-based)
        style: Optional paragraph style name
        font_name: Font family
        font_size: Font size in points
        bold: True/False for bold text
        italic: True/False for italic text
        color: RGB color as hex string
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        body = doc.element.body

        # Get the paragraph at the specified index to insert before it
        paragraphs = doc.paragraphs
        if paragraph_index >= len(paragraphs):
            # Append at the end
            para = doc.add_paragraph(text)
        else:
            # Insert before the paragraph at the specified index
            target_para = paragraphs[paragraph_index]
            new_p = OxmlElement('w:p')
            target_para._element.addprevious(new_p)
            para = doc.paragraphs[paragraph_index]

        # Set text
        para.text = text

        # Apply style
        if style:
            try:
                para.style = doc.styles[style]
            except KeyError:
                pass

        # Apply formatting
        if any([font_name, font_size, bold is not None, italic is not None, color]):
            for run in para.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                if color:
                    color_hex = color.lstrip('#')
                    run.font.color.rgb = RGBColor.from_string(color_hex)

        doc.save(filename)
        return f"Paragraph inserted at index {paragraph_index} in {filename}"
    except Exception as e:
        return f"Failed to insert paragraph: {str(e)}"


async def move_section(filename: str, start_heading_text: str, target_heading_text: str) -> str:
    """Move a section (from start_heading to the next heading of same level) to before target_heading.

    Args:
        filename: Path to the Word document
        start_heading_text: Text of the heading that starts the section to move
        target_heading_text: Text of the heading before which to insert the section
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        paragraphs = doc.paragraphs

        # Find the start heading
        start_idx = None
        start_level = None
        for i, p in enumerate(paragraphs):
            if p.style.name.startswith('Heading') and start_heading_text.lower() in p.text.lower():
                start_idx = i
                start_level = p.style.name
                break

        if start_idx is None:
            return f"Heading '{start_heading_text}' not found"

        # Find the end of the section (next heading of same or higher level)
        end_idx = len(paragraphs)
        for i in range(start_idx + 1, len(paragraphs)):
            p = paragraphs[i]
            if p.style.name == start_level:
                end_idx = i
                break
            # Also check if it's a higher level heading (lower number)
            if p.style.name.startswith('Heading'):
                try:
                    current_level = int(p.style.name.split()[-1])
                    target_level = int(start_level.split()[-1])
                    if current_level <= target_level:
                        end_idx = i
                        break
                except (ValueError, IndexError):
                    pass

        # Find the target heading
        target_idx = None
        for i, p in enumerate(paragraphs):
            if p.style.name.startswith('Heading') and target_heading_text.lower() in p.text.lower():
                target_idx = i
                break

        if target_idx is None:
            return f"Target heading '{target_heading_text}' not found"

        # Collect elements to move (paragraphs and tables between start_idx and end_idx)
        body = doc.element.body
        elements_to_move = []

        # Get all body elements and find the ones between start and end
        para_idx = 0
        for elem in list(body):
            if elem.tag == qn('w:p'):
                if start_idx <= para_idx < end_idx:
                    elements_to_move.append(elem)
                para_idx += 1
            elif elem.tag == qn('w:tbl'):
                # Tables don't have a paragraph index but are between paragraphs
                # Check if the previous paragraph was in range
                if start_idx <= para_idx < end_idx:
                    elements_to_move.append(elem)

        if not elements_to_move:
            return f"No content found to move from section '{start_heading_text}'"

        # Find the target element to insert before
        target_elem = None
        para_idx = 0
        for elem in list(body):
            if elem.tag == qn('w:p'):
                if para_idx == target_idx:
                    target_elem = elem
                    break
                para_idx += 1

        if target_elem is None:
            return f"Target element not found"

        # Move elements
        for elem in elements_to_move:
            body.remove(elem)
            target_elem.addprevious(elem)

        doc.save(filename)
        return f"Section '{start_heading_text}' moved before '{target_heading_text}' in {filename}"
    except Exception as e:
        return f"Failed to move section: {str(e)}"


async def delete_table(filename: str, table_index: int) -> str:
    """Delete a table from the document by its index.

    Args:
        filename: Path to the Word document
        table_index: Index of the table to delete (0-based)
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Table index {table_index} out of range (0-{len(doc.tables)-1})"

        table = doc.tables[table_index]
        table._element.getparent().remove(table._element)

        doc.save(filename)
        return f"Table {table_index} deleted from {filename}"
    except Exception as e:
        return f"Failed to delete table: {str(e)}"


async def insert_table_at_position(filename: str, headers: List[str], data: List[List[str]],
                                     target_text: str = None, target_paragraph_index: int = None,
                                     position: str = 'after') -> str:
    """Insert a formatted table at a specific position in the document.

    Args:
        filename: Path to the Word document
        headers: List of column header strings
        data: 2D list of table data rows
        target_text: Text to find the insertion point (insert before/after this paragraph)
        target_paragraph_index: Paragraph index for insertion point
        position: 'before' or 'after' the target
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."

    try:
        doc = Document(filename)
        paragraphs = doc.paragraphs

        # Find target paragraph
        target_idx = None
        if target_paragraph_index is not None:
            if 0 <= target_paragraph_index < len(paragraphs):
                target_idx = target_paragraph_index
        elif target_text:
            for i, p in enumerate(paragraphs):
                if target_text.lower() in p.text.lower():
                    target_idx = i
                    break

        if target_idx is None:
            # Append at end
            table = doc.add_table(rows=1 + len(data), cols=len(headers))
        else:
            # Create table element and insert at position
            target_para = paragraphs[target_idx]
            # Create a temporary table at the end, then move it
            table = doc.add_table(rows=1 + len(data), cols=len(headers))
            table_elem = table._element
            # Remove from end
            body = doc.element.body
            body.remove(table_elem)
            # Insert at position
            if position == 'before':
                target_para._element.addprevious(table_elem)
            else:
                # Insert after (find next sibling)
                next_sibling = target_para._element.getnext()
                if next_sibling is not None:
                    next_sibling.addprevious(table_elem)
                else:
                    body.append(table_elem)

        # Try to set table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            pass

        # Fill headers
        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = str(h)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Fill data
        for i, row in enumerate(data):
            for j, val in enumerate(row):
                if j < len(headers):
                    cell = table.cell(i + 1, j)
                    cell.text = str(val)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)

        doc.save(filename)
        return f"Table with {len(data)} rows inserted at position {target_idx} ({position}) in {filename}"
    except Exception as e:
        return f"Failed to insert table: {str(e)}"


async def get_paragraph_text(filename: str, paragraph_index: int) -> str:
    """Get the text of a specific paragraph by index.

    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    try:
        doc = Document(filename)
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Paragraph index {paragraph_index} out of range (0-{len(doc.paragraphs)-1})"

        para = doc.paragraphs[paragraph_index]
        return f"Paragraph {paragraph_index} [{para.style.name}]: {para.text}"
    except Exception as e:
        return f"Failed to get paragraph: {str(e)}"
